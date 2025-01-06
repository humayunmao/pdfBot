# Import necessary libraries
from dotenv import load_dotenv
import os
import re
from PyPDF2 import PdfReader
import streamlit as st
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from google.generativeai import GenerativeModel, configure
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import tempfile  # For creating temporary files for downloads
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Pt



# Load environment variables
try:
    load_dotenv()
    GEMINI_API_KEY = "AIzaSyCezBu6X7Pwc2Uii-qtinNQvhjUZCRTQwI"
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not set in the environment variables.")
except Exception as e:
    st.error(f"Error loading environment variables: {e}")
    GEMINI_API_KEY = None

try:
    configure(api_key=GEMINI_API_KEY)
except Exception as e:
    st.error(f"Error configuring Google Gemini API: {e}")

# Set page configuration
try:
    st.set_page_config(
        page_title="LGU PDF_BOT",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
except Exception as e:
    st.error(f"Error setting page configuration: {e}")

# Custom CSS for styling
try:
    st.markdown("""
        <style>
            body {
                background-color: #1B2F1D;
                color: #F2E1C2;
            }
            .main-title {
                font-size: 42px;
                font-weight: bold;
                color: #E1B12C;
                text-align: center;
                margin-bottom: 20px;
            }
            .uploaded-file {
                font-size: 18px;
                color: #C4C4C4;
            }
            .stButton>button {
                background-color: #2A5D37;
                color: white;
                font-size: 16px;
                border-radius: 8px;
            }
            .stButton>button:hover {
                background-color: #3D7F4A;
            }
            .stDownloadButton>button {
                background-color: #C58F22;
                color: white;
                font-size: 16px;
                border-radius: 8px;
            }
            .stDownloadButton>button:hover {
                background-color: #D9A63F;
            }
        </style>
    """, unsafe_allow_html=True)
except Exception as e:
    st.error(f"Error setting custom CSS: {e}")

# Helper functions
def process_text(text):
    try:
        text_splitter = CharacterTextSplitter(
            separator="\n",
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(text)
        embeddings = HuggingFaceEmbeddings()
        knowledge_base = FAISS.from_texts(chunks, embeddings)
        return knowledge_base
    except Exception as e:
        st.error(f"Error processing text: {e}")
        return None

def generate_summary(text):
    try:
        model = GenerativeModel('gemini-pro')
        prompt = f"Provide a concise and accurate summary in a way that highlights the key points and main ideas effectively\n\n{text}"
        summary = model.generate_content(prompt)
        return summary.text
    except Exception as e:
        st.error(f"Error generating summary: {e}")
        return ""

def generate_minute_sheet(text):
    try:
        model = GenerativeModel('gemini-pro')
        prompt = f"Prepare a detailed minutes sheet in tabular format, including the headings: Sr. No., Agenda, Discussion, and Decision. Begin by noting the date, time, venue, topic of the meeting, and a list of all attendees OR Speakers. Then, present the data in the table accordingly. Note: Data must be accurate\n\n{text}"
        
        minute_sheet = model.generate_content(prompt)
        return minute_sheet.text
    except Exception as e:
        st.error(f"Error generating minute sheet: {e}")
        return ""

def create_docx(name, content):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            docx_path = tmp_file.name
            doc = Document()
            title = doc.add_heading(level=1)
            run = title.add_run(name)
            run.font.size = Pt(18)
            run.bold = True
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph)
                    p.style.font.size = Pt(12)
            doc.save(docx_path)
            return docx_path
    except Exception as e:
        st.error(f"Error creating DOCX file: {e}")
        return None

def preprocess_query(query):
    try:
        query = query.lower()
        query = re.sub(r'\W+', ' ', query)  # Remove punctuation
        return query
    except Exception as e:
        st.error(f"Error preprocessing query: {e}")
        return ""

def expand_query(query):
    try:
        model = GenerativeModel('gemini-pro')
        prompt = f"Please rephrase the following query into multiple alternative phrasings:\n\n{query}"
        response = model.generate_content(prompt)
        return response.text.splitlines()
    except Exception as e:
        st.error(f"Error expanding query: {e}")
        return []

def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        text = "".join([paragraph.text + "\n" for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX file: {e}")
        return ""

# Main function
def main():
    try:
        st.markdown('<div class="main-title">📄 LGU PDF_BOT</div>', unsafe_allow_html=True)

        st.sidebar.header("Upload Your Document")
        uploaded_file = st.sidebar.file_uploader("Upload a PDF or DOCX file", type=["pdf", "docx"])

        if uploaded_file:
            # Reset session state if a new file is uploaded
            if "uploaded_file_name" not in st.session_state or st.session_state.uploaded_file_name != uploaded_file.name:
                st.session_state.minute_sheet = None
                st.session_state.summary = None
                st.session_state.knowledge_base = None
                st.session_state.uploaded_file_name = uploaded_file.name

            # Extract and process the document if not already done
            if st.session_state.minute_sheet is None or st.session_state.summary is None:
                if uploaded_file.name.endswith(".pdf"):
                    pdf_reader = PdfReader(uploaded_file)
                    text = "".join([page.extract_text() or "" for page in pdf_reader.pages])
                elif uploaded_file.name.endswith(".docx"):
                    text = extract_text_from_docx(uploaded_file)

                st.session_state.minute_sheet = generate_minute_sheet(text)
                st.session_state.summary = generate_summary(text)
                st.session_state.knowledge_base = process_text(text)

            # Display the results
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("📝 Minute Sheet")
                st.write(st.session_state.minute_sheet)

            with col2:
                st.subheader("🔍 Summary")
                st.write(st.session_state.summary)

            # Provide download options
            st.subheader("📥 Download Results")
            summary_docx_path = create_docx("PDF Summary", st.session_state.summary)
            minute_sheet_docx_path = create_docx("Minute Sheet", st.session_state.minute_sheet)

            col1, col2 = st.columns(2)
            with col1:
                if summary_docx_path:
                    with open(summary_docx_path, "rb") as file:
                        st.download_button(
                            label="Download Summary as DOCX",
                            data=file,
                            file_name="summary.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            with col2:
                if minute_sheet_docx_path:
                    with open(minute_sheet_docx_path, "rb") as file:
                        st.download_button(
                            label="Download Minute Sheet as DOCX",
                            data=file,
                            file_name="minute_sheet.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            # Sidebar search functionality
            st.sidebar.subheader("Search Your Document")
            query = st.sidebar.text_input("Ask a question about the document")

            if query and st.session_state.knowledge_base:
                try:
                    processed_query = preprocess_query(query)
                    expanded_queries = expand_query(processed_query)

                    docs = []
                    for eq in expanded_queries:
                        docs.extend(st.session_state.knowledge_base.similarity_search(eq, k=5))

                    docs = list({doc.page_content: doc for doc in docs}.values())
                    if not docs:
                        st.sidebar.write("No relevant context found in the document. Please rephrase your query.")
                    else:
                        context = "\n\n".join([doc.page_content for doc in docs])
                        prompt = f"Context:\n{context}\n\nQuestion: {query}\nAnswer:"
                        model = GenerativeModel('gemini-pro')
                        response = model.generate_content(prompt)
                        st.sidebar.write("### Answer:")
                        st.sidebar.write(response.text)
                except Exception as e:
                    st.sidebar.error(f"Error processing search query: {e}")
    except Exception as e:
        st.error(f"Unexpected error in the application: {e}")

if __name__ == "__main__":
    main()
